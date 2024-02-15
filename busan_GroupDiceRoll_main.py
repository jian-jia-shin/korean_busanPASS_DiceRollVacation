from PyQt5 import QtWidgets
from PyQt5.QtWidgets import *
from PyQt5.QtCore import *
from PyQt5.QtGui import *

import sys
import time
import busan_GroupDiceRoll_ui

import random
from openpyxl import load_workbook

from docx import Document
from docx.shared import Cm
from datetime import datetime  ###顯示當下時間之模組
#######################設定變數
dice=[0,1,2,3,4,5]
wb = load_workbook("釜山PASS_提供之景點詳細資訊.xlsx")
wbA=wb["Group A"]
wbB=wb["Group B"]
GroupA,GroupB=[],[]
for i in range(2,wbA.max_row+1):
    if i<11:
        GroupA.append("0"+str(i-1)+","+wbA.cell(row=i,column=1).value)
    else:
        GroupA.append(str(i-1)+","+wbA.cell(row=i,column=1).value)
for i in range(2,wbB.max_row+1):
    if i < 11:
        GroupB.append("0"+str(i-1)+","+wbB.cell(row=i,column=1).value)
    else:
        GroupB.append(str(i - 1) + "," + wbB.cell(row=i, column=1).value)
roll_choose_groupA = []
roll_choose_groupB = []
###########################
class myMainWindow(QMainWindow, busan_GroupDiceRoll_ui.Ui_MainWindow):

    def __init__(self):
         super().__init__()
         self.setupUi(self)
         self.pushButton_BIG3.clicked.connect(self.Click_BIG3)
         self.pushButton_BIG5.clicked.connect(self.Click_BIG5)
         self.pushButton_GroupA_roll.setDisabled(True)
         self.pushButton_GroupB_roll.setDisabled(True)
         self.pushButton_GroupA_roll.clicked.connect(self.Click_GroupA_roll)
         self.pushButton_GroupB_roll.clicked.connect(self.Click_GroupB_roll)
         self.pushButton_reset.clicked.connect(self.Click_reset)
         self.pushButton_quit.clicked.connect(self.Click_quit)
         self.listWidget_fullGroupA.addItems(GroupA)
         self.listWidget_fullGroupB.addItems(GroupB)
         self.GroupA_result.setText("")
         self.GroupB_result.setText("")
         self.infomation.setText("")
         self.pushButton_make_word.setDisabled(True)
         self.pushButton_make_word.clicked.connect(self.Make_word)
    def Click_BIG3(self):
        self.pushButton_BIG3.setDisabled(True)
        self.pushButton_BIG5.setDisabled(True)
        time.sleep(3)
        self.pushButton_GroupA_roll.setDisabled(False)
        self.pushButton_GroupB_roll.setDisabled(False)
        self.you_chose_BIG.setText("目前選擇:BIG3")
        self.infomation.setText("目前選擇:BIG3")
        time.sleep(2)
    def Click_BIG5(self):
        self.pushButton_BIG3.setDisabled(True)
        self.pushButton_BIG5.setDisabled(True)
        time.sleep(3)
        self.pushButton_GroupA_roll.setDisabled(False)
        self.pushButton_GroupB_roll.setDisabled(False)
        self.you_chose_BIG.setText("目前選擇:BIG5")
        self.infomation.setText("目前選擇:BIG5")
        time.sleep(2)
    def Click_GroupA_roll(self):
        #time.sleep(1)##緩衝載入時間，避免速度過快而跳出
        choose_BIG = self.you_chose_BIG.text()
        if choose_BIG=="目前選擇:BIG3":
            choseA_limit=1
        if choose_BIG=="目前選擇:BIG5":
            choseA_limit=2

        if len(roll_choose_groupA) != choseA_limit:
            a1, a2, a3 = dice[random.randint(0, 5)], dice[random.randint(0, 5)], dice[random.randint(0, 5)]
            self.dice_view_a1.setText(str(a1))
            self.dice_view_a2.setText(str(a2))
            self.dice_view_a3.setText(str(a3))
            total = a1 + a2 + a3
            if total!=0 or total<=12:
                if GroupA[total-1] not in roll_choose_groupA:
                    roll_choose_groupA.append(GroupA[total-1])
                    print(roll_choose_groupA)
                    self.GroupA_result.setText(f"你擲到{GroupA[total-1]}。")
                    self.listWidget_choosenGroupA.addItem(GroupA[total-1])
                    if len(roll_choose_groupA) == choseA_limit:
                        self.pushButton_GroupA_roll.setDisabled(True)
                        #print(self.listWidget_choosenGroupA.count())
                        self.infomation.setText("GroupA骰選景點已達上限")
                        #time.sleep(1)
                        self.bigN_finish()
                else:
                    self.GroupA_result.setText("有重複景點，再擲一次")
            else:
                self.GroupA_result.setText("小於1，大於12，再擲一次")
    def Click_GroupB_roll(self):
        #time.sleep(1)  ##緩衝載入時間，避免速度過快而跳出
        choose_BIG = self.you_chose_BIG.text()
        if choose_BIG == "目前選擇:BIG3":
            choseB_limit = 2
        if choose_BIG == "目前選擇:BIG5":
            choseB_limit = 3

        if len(roll_choose_groupB) != choseB_limit:
            b1, b2, b3 = dice[random.randint(0, 5)], dice[random.randint(0, 5)], dice[random.randint(0, 5)]
            b4, b5, b6 = dice[random.randint(0, 5)], dice[random.randint(0, 5)], dice[random.randint(0, 5)]
            self.dice_view_b1.setText(str(b1))
            self.dice_view_b2.setText(str(b2))
            self.dice_view_b3.setText(str(b3))
            self.dice_view_b4.setText(str(b4))
            self.dice_view_b5.setText(str(b5))
            self.dice_view_b6.setText(str(b6))
            total = b1 + b2 + b3 + b4 + b5 + b6
            if total != 0 or total <= 22:
                if GroupB[total - 1] not in roll_choose_groupB:
                    roll_choose_groupB.append(GroupB[total - 1])
                    print(roll_choose_groupB)
                    self.GroupB_result.setText(f"你擲到{GroupB[total - 1]}。")
                    self.listWidget_choosenGroupB.addItem(GroupB[total - 1])
                    if len(roll_choose_groupB) == choseB_limit:
                        self.pushButton_GroupB_roll.setDisabled(True)
                        self.infomation.setText("GroupB骰選景點已達上限")
                        #time.sleep(1)
                        self.bigN_finish()
                else:
                    self.GroupB_result.setText("有重複景點，再擲一次")
            else:
                self.GroupB_result.setText("小於1，大於22，再擲一次")
    def bigN_finish(self):
        if self.you_chose_BIG.text() == "目前選擇:BIG3":
            if self.listWidget_choosenGroupA.count() == 1 and self.listWidget_choosenGroupB.count() == 2:
                print(roll_choose_groupA[0][0:2])
                self.infomation.setText("達成BIG3骰選景點，確定的話可以點\"產生結果文件\"")
                self.pushButton_make_word.setDisabled(False)
        if self.you_chose_BIG.text() == "目前選擇:BIG5":
            if self.listWidget_choosenGroupA.count() == 2 and self.listWidget_choosenGroupB.count() == 3:
                print(roll_choose_groupA[0][0:2])
                self.infomation.setText("達成BIG5骰選景點，確定的話可以點\"產生結果文件\"")
                self.pushButton_make_word.setDisabled(False)
    def Click_reset(self):
        roll_choose_groupA.clear()
        roll_choose_groupB.clear()
        self.listWidget_choosenGroupA.clear()
        self.listWidget_choosenGroupB.clear()
        self.pushButton_BIG3.setDisabled(False)
        self.pushButton_BIG5.setDisabled(False)
        self.pushButton_GroupA_roll.setDisabled(True)
        self.pushButton_GroupB_roll.setDisabled(True)
        self.GroupA_result.setText("")
        self.GroupB_result.setText("")
        self.infomation.setText("已重新設定")
        self.you_chose_BIG.setText("目前選擇:")
        self.pushButton_make_word.setDisabled(True)
    def Make_word(self):
        document = Document()
        doc_section = document.sections
        for section in doc_section:
            # 設定分節頁面左邊界1公分
            section.left_margin = Cm(1)
            # 設定分節頁面右邊界1公分
            section.right_margin = Cm(1)
            # 設定分節頁面上邊界1公分
            section.top_margin = Cm(1)
            # 設定分節頁面下邊界1公分
            section.bottom_margin = Cm(1)
        if self.you_chose_BIG.text()=="目前選擇:BIG3":
            BIGN="釜山PASS_BIG3骰選景點"
            document.add_heading('釜山PASS_BIG3骰選景點結果', 0)
        if self.you_chose_BIG.text()=="目前選擇:BIG5":
            BIGN = "釜山PASS_BIG5骰選景點"
            document.add_heading('釜山PASS_BIG5骰選景點結果', 0)
        def writinng(GroupLimit,getRollChoose,getwb,GroupTitle):
            i=0
            document.add_paragraph(GroupTitle)
            while i<GroupLimit:
                getNum=int(getRollChoose[i][0:2])+1##取得景點編號
                document.add_heading(getwb.cell(row=getNum,column=1).value, level=1)
                for se in range(3,getwb.max_column+1):
                    document.add_paragraph(getwb.cell(row=1,column=se).value, style='Intense Quote')
                    document.add_paragraph(getwb.cell(row=getNum,column=se).value)
                document.add_page_break()
                i+=1
        writinng(len(roll_choose_groupA),roll_choose_groupA,wbA,"GroupA骰選景點")
        writinng(len(roll_choose_groupB), roll_choose_groupB, wbB,"GroupB骰選景點")
        now_time=datetime.now().strftime("%Y%m%d_%H%M%S")
        document.save(now_time+"_"+BIGN+'結果文件'+".docx")
        time.sleep(2)
        self.infomation.setText("完成產出，請點擊重新設定")
        self.pushButton_make_word.setText("產生\n結果文件")
        self.pushButton_make_word.setDisabled(True)

    def Click_quit(self):
        sys.exit(app.exec_())


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    window = myMainWindow()
    window.show()
    sys.exit(app.exec_())